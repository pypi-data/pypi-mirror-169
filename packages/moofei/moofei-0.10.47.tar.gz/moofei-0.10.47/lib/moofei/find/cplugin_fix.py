#!/usr/bin/python
# -*- coding: utf-8 -*-
# editor: mufei(ypdh@qq.com tel:+086 15712150708)
'''
Mufei _ __ ___   ___   ___  / _| ___(_)
| '_ ` _ \ / _ \ / _ \| |_ / _ \ |
| | | | | | (_) | (_) |  _|  __/ |
|_| |_| |_|\___/ \___/|_|  \___|_|
'''

__all__ = ['_StringIO','StringIO','BytesIO', 'cPlugin_Fix', 'is_ascii']

import warnings
import os,sys,time
import copy
try:
    from io import StringIO ## for Python 3
    from io import BytesIO
except ImportError:
    from StringIO import StringIO ## for Python 2
    from cStringIO import StringIO as BytesIO    
   
from moofei._find import _py, _strtypes, _find_func, _search_func, _get_detect_bytes, _fcall, _get_newline_bytes
from moofei._valid import VALID

def is_ascii(s):
    '''is_ascii(byte) -> bool
    '''
    if (ord(s) in range(32, 122) ) or (ord(s) in range(1039, 1104))\
    or (ord(s) in (9,10,13)):
        return True
    return False
    
class FindError(BaseException):
    def __init__(self, msg=None):
        self.msg = msg
    def __str__(self):
        if self.msg is None:
            msg = ''
        elif not isinstance(self.msg, str):
            msg = str(self.msg) 
        else:
            msg = self.msg
        return  msg  
        
class _StringIO:
    def __init__(self, text=b""):
        self.text = text
        self.index = 0        
    def write(self, text):
        #if text == u'\xe9':return 
        try: 
            text = bytes(text)
        except:
            text = "" #text.encode('utf-8')   
        if text: self.text += text    
    def seek(self, index):self.index = index    
    def getvalue(self):return self.text.decode('utf-8')   
    def close(self): pass 
    def __enter__(self):return self
    def __exit__(self, exc_type, exc_val, exc_tb):return False

class cPlugin_Fix:
    def __init__(self, fpath, words=None, wcase=None, fp=None,
                       **awgs):
        '''
        CSV格式文件：分隔符 ',' comma逗号
        PSV格式文件：分隔符'|',管道文件pipe-delimited text files
        TSV 格式文件：分隔符' '  tab空格
        '''    
        self.fpath = fpath
        self.words = words
        self.wcase = wcase 
        self.fp = fp        
        self.__awgs = awgs
        self.frombox = self.__awgs.get('frombox')
        self.sleep = self.__awgs.get('sleep')
        self.is_fix_image = self.__awgs.get('is_fix_image')
        self.__boxfname= None
        
    def is_enable(self):
        return True
        
    def result(self):
        text = self.all_texts()
        return self.find(text)
    
    def find_chunk(self, boxfname, chunk):
        if self.sleep: self.sleep(0)
        if _py[0]==2 and isinstance(boxfname, type(u'')) : #py2.x unicode
            boxfname = boxfname.encode('gbk') #此处GBK应该来自--coding 有变动
        if getattr(chunk, 'read', None): 
            chunk.seek(0)
            fpArgs=(chunk,) 
        else:
            fpArgs=(BytesIO(chunk),)
        awgs = copy.deepcopy(self.__awgs)
        awgs['frombox'] = boxfname
        #awgs['wcall'] = None
        return _fcall(self.fpath, words=self.words, wcase=self.wcase, #frombox=boxfname, 
                      fpArgs=fpArgs,
                      **awgs)
               
    def find_chunk_txt(self, boxfname, chunk, coding=None):
        #遇到boxfname(docx,pdf等还需要切换搜索)
        if self.sleep: self.sleep(0)
        _words, _coding =  _get_detect_bytes(self.words, coding=coding, chunk=chunk[:100*1024], fpath=boxfname)         
        _chunk = chunk.lower()  if self.wcase==1 else chunk
        if isinstance(self.words, _strtypes):
            func = _find_func
            if self.wcase==1 : _words = _words.lower()    
        else:
            func = _search_func             
        return func(_chunk, _words, fpath=self.fpath, newline=_get_newline_bytes(_coding),
                    encoding=_coding, **self.__awgs
               )
        
    def find(self, text, coding=None):
        encoding = 'str'
        if self.sleep: self.sleep(0)
        if not text:
            return []
        elif isinstance(text, bytes):
            return self.find_chunk_txt(self.__awgs.get('frombox', self.fpath), text, coding=coding)   
        
        elif isinstance(self.words, _strtypes):
            func = _find_func
            words = self.words
            if self.wcase==1 : self.words = self.words.lower()    
        else:
            func = _search_func
            words = self.words.pattern
        if self.wcase==1 : 
            text = text.lower() 
        if _py[0] == 2 and type(words)!=type(text):       
            if isinstance(self.words, str):
                #text = text.encode('gbk')
                try:
                    self.words = self.words.decode('utf-8')
                except:
                    self.words = self.words.decode('gbk')
            elif isinstance(text, unicode):
                #text = VALID.isStr(text)[1]
                try:
                    text = text.encode('gbk')
                    encoding = 'gbk'
                except:
                    text = text.encode('utf-8')
                    encoding = 'utf-8'
        return func(text, self.words, fpath=self.fpath, 
                    encoding=encoding, **self.__awgs
                  )         
    
    def get_tempfile_fix_path(self, suffix):
        import tempfile
        if suffix[0]!='.': suffix = '.'+suffix.split('.')[-1]
        fp = tempfile.NamedTemporaryFile(delete=False,suffix=suffix) # 
        return fp.name
                   
    def get_tempfile_or_path(self, force=False, suffix=None):
        import tempfile
        if self.__awgs.get('frombox') or force:
            if suffix:
                if suffix is True: suffix=self.__class__.__name__[10:].lower().replace('_','.')
                if suffix.split('.')[-1] in ('fix','image','audio'):
                    suffix = '.'+self.__awgs.get('frombox').split('.')[-1]
                fp = tempfile.NamedTemporaryFile(delete=False,suffix=suffix) # 
            else:
                fp = tempfile.NamedTemporaryFile(delete=False) #
            self.__boxfname = fp.name
            if self.fp:
                fp.write(self.fp.read())
            else:
                fp.write(open(self.fpath, 'rb').read())
            fp.close()
            return self.__boxfname
        else:
            return self.fpath
    
    def image2txt(self, fpath,  fp, frombox=None):
        try:
            from .plugin_fix_image import cPlugin_Fix_Image  
        except:
            from moofei.find.plugin_fix_image import cPlugin_Fix_Image  
        o = cPlugin_Fix_Image(fpath, fp=fp, frombox=frombox)
        if self.sleep: self.sleep(0)
        return o.all_texts()
        
    def word2txt(self, suffix=None):
        #from threading import Thread 
        from moofei.thread import Thread
        t = Thread(target=self.__word2txt, suffix=suffix)
        t.start()
        t.join(30) #最长等待30秒
        if t.isAlive(): #防止Word文件被锁定进程一直等待
            try:
                t._stop()
            except:
                print('Stoped ...', self.fpath, self.__boxfname)
            raise FindError((self.words, t.is_error, t.fncall_result))    
        return t.fncall_result
                 
    def __word2txt(self, suffix=None):
        import pythoncom; pythoncom.CoInitialize()
        from win32com.client import Dispatch, DispatchEx
        import win32com
        wc=win32com.client.constants
        
        #word = win32com.client.GetActiveObject('Word.Application')
        #word = win32com.client.GetObject(Class = "Word.Application")
        #word = win32com.client.gencache.EnsureDispatch('Word.Application.8')
        
        #word = Dispatch('Word.Application')     # 打开word应用程序
        try:
            word = DispatchEx('Word.Application') # 启动独立的进程
        except:
            os.popen('taskkill /IM winword.exe /F').read()
            word = DispatchEx('kwp.application')
        #word = DispatchEx('kwps.application')
        #except pythoncom.com_error:
            #print "Starting Word 7 for dynamic test"
            #word = win32com.client.Dispatch("Word.Basic")
            
        #1.安装软件capicom_dc_sdk.msi
        #2.打到capicom.dll，并复制到SysWOW64下： 
        #3.执行以下命令：cd C:\Windows\SysWOW64 C:\Windows\SysWOW64>regsvr32.exe capicom.dll 
        #4.安全性问题: 在Microsoft Word文档的DCOM里的identity选项里，设置成用管理员帐户执行即可
        
        word.Visible = 0        # 后台运行,不显示
        word.DisplayAlerts = 0  # 不警告
        Ls = []
        
        if self.fpath.startswith('/') or ':' in self.fpath or self.__awgs.get('frombox'):
            fpath = self.get_tempfile_or_path(suffix=suffix or True)
        else:
            fpath = os.path.join(os.getcwd(), self.fpath)    
        doc = word.Documents.Open(FileName=fpath, ReadOnly=True, Encoding='gbk')
        #doc = word.Documents.Open(self.fp or self.fpath, Encoding='gbk')
        try:
            if not doc:
                if word.ActiveDocument:
                    doc = word.ActiveDocument
                    #return doc.Content.Text
                else:
                    return None #需要打开word文档用户和Python运行用户一致 
            if self.sleep: self.sleep(0)
            if not getattr(doc, 'paragraphs', None):
                return doc.Content.Text
            for para in doc.paragraphs:
                Ls.append(para.Range.Text)
            
            for t in doc.Tables:
                s = ''
                for row in t.Rows:
                    for cell in row.Cells:
                        #s+=' '
                        s+=cell.Range.Text
                Ls.append(s)
        
            doc.Close()
            #doc.Close(wc.wdDoNotSaveChanges)
        finally:
            word.Quit()
        del doc
        del word
        text = '\n'.join(Ls)
        return text
        
    def docx2txt(self, suffix=None):
        from docx import Document
        if self.fpath.startswith('/') or ':' in self.fpath or self.__awgs.get('frombox'):
            fpath = self.get_tempfile_or_path(suffix=suffix or True)
        else:
            fpath = os.path.join(os.getcwd(), self.fpath) 
        document = Document(self.fp or fpath)
        if self.sleep: self.sleep(0)
        Ls = []
        for paragraph in document.paragraphs:
            Ls.append(paragraph.text)
        return '\n'.join(Ls)
    
    @classmethod    
    def docx2pdf(cls, src, dst):
        '''
        word = client.DispatchEx("Word.Application")
        worddoc = word.Documents.Open(doc_name,ReadOnly = 1)
        
        #添加宏
        macro = worddoc.VBProject.VBComponents.Add(1)
        macro.CodeModule.AddFromString(macrocode)
        
        #执行已有宏文件
        #注意，当宏名是唯一，不需要写模块名，但如果模块名和宏名重复，需要补全
        marco=worddoc.macro('m.vba_marco')   
        marco()
        #或 worddoc.Application.Run(VBA_name, args) ##args 为宏参数
        time.sleep(50)   #等待macro运行完毕，具体等待市场视宏运行的时间

        worddoc.SaveAs(pdf_name, FileFormat = 17)
        worddoc.Close()
        '''
        if not os.path.exists(src):
            print(src + "不存在，无法继续！")
            return False
        os.system('taskkill /im wps.exe')
        # 如果文件存在就删除
        if os.path.exists(dst):
            os.remove(dst)
        import win32com.client    
        o = win32com.client.Dispatch("Kwps.Application")
        #if self.sleep: self.sleep(0)
        o.Visible = False
        doc = o.Documents.Open(src)
        doc.ExportAsFixedFormat(dst, 17)
        doc.Close()
        o.Quit()
        if os.path.exists(dst):
            return True
        else:
            return False
            
    def watermark_word(self, content, fpath, dist=None, **awgs):
        '水印'
        if isinstance(fpath, (list,tuple)):
            suffix = '.pdf'
            if dist and isinstance(dist,(list,tuple)): dist=None
        else:
            suffix = fpath
        if not dist : dist = self.get_tempfile_fix_path(suffix=fpath)    
        if suffix.endswith('.pdf'):
            from moofei._pdf import _Pdf
            opath = _Pdf.watermark_word(content, fpath, out_path=dist, **awgs)
        else:
            from moofei._image import _Image
            opath = _Image.watermark_word(content, fpath, out_path=dist, **awgs)    
        if self.sleep: self.sleep(0)
        return dist 
            
    def __del__(self):
        if self.__boxfname:
            os.remove(self.__boxfname)
            self.__boxfname = None
            
    @classmethod
    def _main(cls):
        #(eindex+start,eindex+end, i+sindex, countn)
        fname = input('Entry path: ')
        Search = input('Search word: ')
        this=cls(fname, Search)
        if Search:
            print(this.result())
        else:
            print(this.all_texts())
            
        
        
            
        

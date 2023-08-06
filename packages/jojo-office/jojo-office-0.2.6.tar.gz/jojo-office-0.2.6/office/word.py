#
# Word : auto creation of the .docx file
#
# Example:
# ::
#     import office
#
#     # create file output.docx using template1.docx
#     doc = office.open_file("output.docx", template="template1.docx")
#
#     # create document content by fill the data from Excel file datafile.xlsx, and save.
#     doc.fill('datafile.xlsx').save()
#
#     # save to pdf with watermark (works on Windows with Microsoft PowerPoint)
#     doc.save('final.pdf', watermark="CONFIDENTIAL")
#
#

import os
import time
import copy
import docx
from types import MethodType
from .base import BaseFile
from .util import Util
from .variable import Variable, VarData
from .win32 import Win32
from .pdf import PDF


class WordConsts:
    # https://docs.microsoft.com/en-us/office/vba/api/word.wdsaveformat
    wdFormatDocument = 0   # Microsoft Office Word 97 - 2003 binary file format.
    wdFormatTemplate = 1  # Word template format.
    wdFormatText = 2  # Microsoft Windows text format.
    wdFormatRTF = 6  # Rich text format (RTF).
    wdFormatHTML = 8  # Standard HTML format.
    wdFormatWebArchive = 9  # Web archive format.
    wdFormatFilteredHTML = 10  # Filtered HTML format.
    wdFormatXMLTemplate = 14  # XML template format.
    wdFormatPDF = 17  # PDF format.


# Word .docx Document
class Word(BaseFile):
    """ Word .docx Document"""

    class Paragraphs:
        """ collection of paragraph """
        def __init__(self, word):
            self.word = word
            # noinspection PyProtectedMember
            self.document = word._obj

        def __len__(self):
            return len(self.document.paragraphs)

        def _add_method(self, paragraph):
            """ add method to page object """
            def move_to(this, new_index):
                this.word.move_paragraph(this.get_index(), new_index)

            def get_index(this):
                return this.element.getparent().index(this.element)

            def delete(this):
                this.word.delete_paragraph(this.get_index())

            # noinspection PyProtectedMember
            paragraph.element = paragraph._element
            paragraph.word = self.word
            paragraph.move_to = MethodType(move_to, paragraph)
            paragraph.delete = MethodType(delete, paragraph)
            paragraph.get_index = MethodType(get_index, paragraph)
            return paragraph

        def __getitem__(self, item):
            if 0 <= item < len(self.document.paragraphs):
                p = self.document.paragraphs[item]
                return self._add_method(p)
            raise IndexError('Index %s out of bound' % repr(item))

    def _open(self, filename, **kwargs):
        self._obj = docx.Document(filename)
        self.filename = filename
        return True

    def _close(self):
        self._obj = None
        self.filename = None
        return True

    def _new_file(self, filename, **kwargs):
        self._obj = docx.Document()
        if filename:
            self.save(filename)

    def _save_file(self, filename, **kwargs):
        formats = {
            '.txt': WordConsts.wdFormatRTF,
            '.rtf': WordConsts.wdFormatText,
            '.dot': WordConsts.wdFormatTemplate,
            '.dotx': WordConsts.wdFormatXMLTemplate,
            '.doc': WordConsts.wdFormatDocument,
            '.pdf': WordConsts.wdFormatPDF,
            '.html': WordConsts.wdFormatHTML,
            '.htm': WordConsts.wdFormatHTML,
            '.mht': WordConsts.wdFormatWebArchive,
            '.mhtml': WordConsts.wdFormatWebArchive,
        }

        file_ext = Util.file_ext(filename)

        if file_ext == '.docx':
            self._obj.save(filename, **kwargs)
            self.filename = filename
            return True

        elif file_ext in formats:
            app, document = Win32.open_word(self.filename, False)
            if not Util.is_absolute_path(filename):
                filename = os.path.join(os.getcwd(), filename)

            password = kwargs.pop('password', None)
            watermark = kwargs.pop('watermark', None)

            document.SaveAs(filename, formats[file_ext], **kwargs)
            Win32.close(app, document)

            if (password or watermark) and Util.file_ext(filename) == '.pdf':
                PDF(filename).watermark(watermark).save(password=password)

        else:
            return super()._save_file(filename, **kwargs)

    def _fill_data(self, data):
        """
        fill the data, replace variable tags and create file content.

        :param data:  dictionary or Workbook object or DataFrame object
        :return: self
        """
        row = 0
        level = 0

        # 逐个段落处理
        i = 0
        while i < len(self._obj.paragraphs):
            paragraph = self.paragraphs[i]
            i += 1
            result, level = self._fill_paragraph(paragraph, data, row, level, ignore=False)
            if hasattr(result, 'element'):  # 发现 @repeat 时，result 是 Paragraph 对象
                # 处理复制段落
                i = self._repeat_paragraph(i - 1, data)
                level = 0

        return self

    def _repeat_paragraph(self, index, var_data):
        """ create repeat paragraphs """
        # 1，发现 @repeat.
        # 2, 找到相关段落、变量名列表
        # 3，找到变量的最大数量N
        # 4，复制段落（备份）N 次
        # 5，处理复制段落(忽略@repeat)

        # 找到需要重复的段落、变量名列表
        indexes, var_names = self._find_repeat_paragraphs(index)
        # 找到变量的最大行数量 repeat_count
        repeat_count = 0
        for v in var_names:
            count = var_data.var_length(v)
            count = 1 if not isinstance(count, int) else count
            if count > repeat_count:
                repeat_count = count

        if repeat_count == 0:
            # 如果repeat_count=0，删除段落
            for _ in indexes:
                self.delete_paragraph(index)
            return index
        else:
            # 如果repeat_count>0， 复制段落 repeat_count 次
            copied = 0
            for n in range(repeat_count - 1):
                for idx in indexes:
                    new_index = index + len(indexes) + copied
                    self.copy_paragraph(idx, new_index)
                    copied += 1

            # 处理重复段落(忽略@repeat)
            for row in range(1, repeat_count + 1):
                level = 0
                for line in range(len(indexes)):
                    i = index + (row - 1) * len(indexes) + line
                    paragraph = self.paragraphs[i]
                    result, level = self._fill_paragraph(paragraph, var_data, row, level, ignore=True)
            # 返回重复段落后的段落
            index = index + copied + len(indexes)
        return index

    def _insert_object(self, paragraph, filename):
        pass

    def _fill_paragraph(self, paragraph, data: VarData, row, level, ignore=False):
        """ fill the data into the paragraph """

        def process(var):
            """ process a var """
            value = data.get_var(var, row)
            value = str(value) if value is not None else ''
            if value.startswith(Variable.OBJECT_PREFIX):  # 如果 value 是对象
                # noinspection PyBroadException
                try:
                    # 插入对象
                    self._insert_object(paragraph, value[len(Variable.OBJECT_PREFIX):])
                except Exception:
                    pass
                return ''
            else:
                return value

        varname = ''
        for run in paragraph.runs:
            old_text = run.text
            text = ''
            for idx, c in enumerate(old_text):
                if c == "{":
                    varname += c
                    level += 1
                elif c == "}":
                    varname += c
                    level -= 1
                    if level == 0:
                        if data.has_var(varname):
                            val = process(varname)
                            text += val
                        else:
                            text += varname
                        varname = ''
                else:
                    if level > 0:
                        varname += c
                        if varname == '{@repeat':
                            if ignore:
                                varname = ''
                                level = level - 1
                            else:
                                run.text = text + varname
                                return paragraph, level
                    else:
                        text += c

            run.text = text
        return None, level

    def _find_repeat_paragraphs(self, i):
        """ return (list_paragraph_object, list_varname) """
        varname = ''
        level = 0
        result = []
        var_names = []
        for i in range(i, len(self.paragraphs)):
            paragraph = self.paragraphs[i]
            result.append(i)
            for run in paragraph.runs:
                for idx, c in enumerate(run.text):
                    if c == "{":
                        varname = c
                        level += 1
                    elif c == "}":
                        varname += c
                        level -= 1
                        if Variable.is_var(varname):
                            var_names.append(varname)
                        varname = ''
                        if level == 0:
                            return result, var_names
                    else:
                        if level > 0:
                            varname += c
                        else:
                            pass
        return result, var_names

    @property
    def paragraphs(self):
        """ return paragraphs """
        return Word.Paragraphs(self)

    @staticmethod
    def _find_paragraph(doc_obj, paragraph):
        for index, p in enumerate(doc_obj.paragraphs):
            # noinspection PyProtectedMember
            if p._element == paragraph._element:
                return index
        return -1

    def find_paragraph(self, paragraph):
        return Word._find_paragraph(self._obj, paragraph)

    def copy_paragraph(self, index, new_index=None, dst_doc=None):
        """ copy paragraph """
        if dst_doc is None:
            dst_doc = self._obj
        elif isinstance(dst_doc, Word):
            dst_doc = dst_doc._obj

        # create paragraph
        src = self._obj.paragraphs[index]
        dst = dst_doc.add_paragraph('', style=src.style)
        index = Word._find_paragraph(dst_doc, dst)
        # copy runs
        for run in src.runs:
            output_run = dst.add_run(run.text, style=run.style)
            output_run.bold = run.bold  # Run's bold data
            output_run.italic = run.italic  # Run's italic data
            output_run.underline = run.underline  # Run's underline data
            output_run.font.color.rgb = run.font.color.rgb  # Run's color data
        # copy paragraph_format
        src_fmt = copy.deepcopy(src.paragraph_format.element)
        dst_fmt = dst.paragraph_format.element
        dst_fmt.getparent().replace(dst_fmt, src_fmt)
        # move to new place if needed
        if new_index:
            return self.move_paragraph(index, new_index)
        else:
            return index

    def delete_paragraph(self, index):
        """ delete paragraph """
        # noinspection PyProtectedMember
        p = self.paragraphs[index]._element
        p.getparent().remove(p)
        p._p = p._element = None

    def move_paragraph(self, old_index, new_index):
        """ move paragraph """
        if old_index == new_index and 0 < old_index < len(self.paragraphs):
            return old_index

        # noinspection PyProtectedMember
        p = self.paragraphs[old_index]._element
        parent = p.getparent()
        parent.remove(p)
        parent.insert(new_index, p)
        index = parent.index(p)
        return index

    def print(self, **kwargs):
        """ print file on default printer """
        print("printing", self.filename, '...')

        app, document = Win32.open_word(self._absolute_filename, False)
        document.PrintOut(**kwargs)
        time.sleep(1)
        Win32.close(app, document)

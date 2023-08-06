
# 安装python-docx包：pip install python-docx -i http://pypi.douban.com/simple --trusted-host pypi.douban.com
import os
import yaml
import time
import pathlib
import pytest
from pyecharts.charts import Bar
from pyecharts import options as opts
from pyecharts.charts import Line
from docx.shared import Inches
from docx import Document
from pyecharts.charts import Pie
from pyecharts.render import make_snapshot
from snapshot_selenium import snapshot
from generate_report.get_db import get_db
from generate_report.chrome_driver import *

file_path = os.path.dirname(os.path.abspath(__file__))
path = f'{file_path}/project_data.yml'
with open(path, encoding="utf-8") as f:
    config = yaml.full_load(f)

db = get_db("db_zentao")

project_li = []
project_data = config.get('info')
project_li.append(project_data)


class TestReport:

    @pytest.mark.parametrize('project', project_li)
    def test_connect_db(self, project):
        project_name = project.get('project_name')
        writer = project.get('writer')
        expect_di = project.get('di')
        generate_path = project.get('path')
        # 获取测试报告存放的路径
        report_path = f'{generate_path}/{project_name}测试报告.docx'
        # 查看结果文件夹是否存在，若不存在，则创建
        if os.path.exists(f'{file_path}/report_res/{project_name}'):
            pass
        else:
            os.makedirs(f'{file_path}/report_res/{project_name}', exist_ok=True)
        # 查看结果文件是否存在，若不存在，则创建
        if os.path.exists(report_path):
            os.remove(report_path)
            pathlib.Path(report_path).touch()
        else:
            pathlib.Path(report_path).touch()
        # 写入项目名称、撰写人、撰写日期
        self.write_string(report_path, '{{project_name}}', project_name)
        self.write_string(report_path, '{{writer}}', writer)
        now_time = time.strftime('%Y-%m-%d  %H:%M', time.localtime())
        self.write_string(report_path, '{{write_time}}', now_time)
        # 根据项目名称获取项目id
        project_id = db.select_data(f"select id from zt_project where `name`='{project_name}';")
        # 计算bug数量
        pid = int(project_id[0])
        result = db.select_data(
            f"select count(*) from zt_bug where project={pid} and deleted='0' and assignedTo='closed';")
        bug_num = str(result[0])
        self.write_string(report_path, '{{bug_number}}', bug_num)
        checkChromeDriverUpdate()
        # 按模块分类的bug--表格
        module_result = db.select_data_list(f"""
                    select r.`name`,r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                    from (select a.*,c.`name` from (select module,count(*) as count from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed' group by module)a
                    inner join zt_module as c on a.module=c.id)r, (select count(*) as total from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, module_result, '<<module_table>>')
        # 按模块分类的bug--饼图
        title = '模块bug数量'
        # 绘制饼图
        self.write_pie(project_name, title, module_result)
        image_1 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的饼图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_1, '<<module_image>>')
        # 查找最多bug数量的模块
        max_bug = self.find_max_bug(module_result)
        self.write_string(report_path, '{{max_bug_module}}', max_bug[0])
        # 按严重程度分类的bug--表格
        level_result = db.select_data_list(f"""
                    select (case when r.severity=1 then 'A'
                    when r.severity=2 then 'B'
                    when r.severity=3 then 'C'
                    when r.severity=4 then 'D'
                    when r.severity=5 then 'E'
                    end)level,r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                    from (select severity,count(*) as count from zt_bug where project={pid} and deleted='0'
                     and assignedTo='closed' group by severity)r, (select count(*) as total from zt_bug
                     where project={pid} and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, level_result, '<<bug_level_table>>')
        # 按bug严重程度统计--饼图
        title = '按bug严重程度统计'
        # 绘制饼图
        self.write_pie(project_name, title, level_result)
        image_2 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的饼图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_2, '<<bug_level_image>>')
        # 按解决方案分类的bug--表格
        resolve_result = db.select_data_list(f"""
                    select (case when r.resolution='fixed' then '开发已修复'
                    when r.resolution='external' then '不是问题'
                    when r.resolution='duplicate' then '重复bug'
                    when r.resolution='datafixed' then '数据已修复'
                    when r.resolution='externalReasons' then '外部原因'
                    when r.resolution='notrepro' then '无法重现'
                    when r.resolution='postponed' then '延期'
                    when r.resolution='xxqysx' then '新需求已实现'
                    when r.resolution='ljpcygz' then '理解偏差已更正'
                    when r.resolution='ywycl' then '运维已处理'
                    end)level,r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                    from (select resolution,count(*) as count from zt_bug where project={pid} and deleted='0'
                    and assignedTo='closed' group by resolution)r,(select count(*) as total from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, resolve_result, '<<resolve_table>>')
        # 按bug解决方案统计--饼图
        title = '按bug解决方案统计'
        # 绘制饼图
        self.write_pie(project_name, title, resolve_result)
        image_3 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的饼图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_3, '<<resolve_image>>')
        # 查找数量最多的解决方案
        max_bug_resolve = self.find_max_bug(resolve_result)
        self.write_string(report_path, '{{max_bug_resolve}}', max_bug_resolve[0])
        self.write_string(report_path, '{{fix_proportion}}', max_bug_resolve[-1])
        # 查询延期处理的bug
        # bug数量
        for i in resolve_result:
            if i[0] == '延期':
                num = i[1]
            else:
                num = 0
            self.write_string(report_path, '{{second_bug_number}}', str(num))
        # 延期处理bug的详细信息
        postpone_result = db.select_data_list(f"""select r.id,r.title,'无' as explains,'延期' as advise
                            from (select * from zt_bug
                            where project={pid} and deleted='0' and assignedTo='closed' and resolution='postponed')r,
                            (select count(*) as total from zt_bug
                            where project={pid} and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, postpone_result, '<<postpone_bug_table>>', 1)
        # 按激活次数分类的bug--表格
        activated_result = db.select_data_list(f"""
                    select concat('激活次数：', r.activatedCount, '次'),
                    r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                    from (select activatedCount,count(*) as count from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed' group by activatedCount)r,
                    (select count(*) as total from zt_bug where project={pid}
                    and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, activated_result, '<<activated_count_table>>')
        # 按bug激活次数统计--饼图
        title = '按bug激活次数统计'
        # 绘制饼图
        self.write_pie(project_name, title, activated_result)
        image_4 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的饼图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_4, '<<activated_count_image>>')
        # 查找激活次数为0的bug数量
        for j in activated_result:
            if j[0] == '激活次数：0次':
                zero_num = j[-1]
                self.write_string(report_path, '{{activated_zero}}', zero_num)
        # 按bug类型分类的bug--表格
        type_result = db.select_data_list(f"""
                    select (case when r.type='codeerror' then '代码错误'
                    when r.type='designchange' then '设计变更'
                    when r.type='hxwt' then '红线问题'
                    when r.type='security' then '安全相关'
                    when r.type='xqsj' then '需求设计'
                    when r.type='sjwt' then '数据问题'
                    when r.type='yhjy' then '优化建议'
                    when r.type='install' then '环境部署'
                    when r.type='interface' then '界面优化'
                    when r.type='config' then '性能问题'
                    end)level,r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                    from (select type,count(*) as count from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed' group by type)r,
                    (select count(*) as total from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, type_result, '<<bug_type_table>>')
        # 按bug类型统计--饼图
        title = '按bug类型统计'
        # 绘制饼图
        self.write_pie(project_name, title, type_result)
        image_5 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的饼图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_5, '<<bug_type_image>>')
        # 绘制帕累托图
        title = '帕累托图分析'
        self.get_pareto(project_name, title, type_result)
        image_7 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_7, '<<Pareto_Chart>>')
        # 按开发人员分类的bug--表格
        people_result = db.select_data_list(f"""
                    select r.realname,r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                    from (select a.*,c.realname from (select resolvedBy,count(*) as count from zt_bug
                    where project={pid} and deleted='0' and assignedTo='closed' group by resolvedBy)a
                    inner join zt_user as c on a.resolvedBy=c.account)r,
                    (select count(*) as total from zt_bug 
                    where project={pid} and deleted='0' and assignedTo='closed')b;""")
        self.write_table(report_path, people_result, '<<exploitation_table>>')
        # 按开发人员解决的bug统计--饼图
        title = '每人解决的bug数'
        # 绘制饼图
        self.write_pie(project_name, title, people_result)
        image_6 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的饼图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_6, '<<exploitation_image>>')
        # 按测试人员每天提交的bug--折线图
        submit_result = db.select_data_list(f"""
                        select c.realname,a.date,a.count from 
                        (select openedBy,date_format(openedDate,'%Y-%m-%d') as date,count(*) as count from zt_bug
                        where project={pid} and deleted='0' and assignedTo='closed' group by openedBy,date)a
                        inner join zt_user as c on a.openedBy=c.account;;""")
        title = '测试人员每天提交的bug数'
        # 绘制折线图
        self.write_line(project_name, title, submit_result)
        image_8 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的折线图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_8, '<<test_submit_days>>')
        # 按测试人员每天关闭的bug--折线图
        close_result = db.select_data_list(f"""
                        select c.realname,a.date,a.count from 
                        (select closedBy,date_format(closedDate,'%Y-%m-%d') as date,count(*) as count from zt_bug
                        where project={pid} and deleted='0' and assignedTo='closed' group by closedBy,date)a
                        inner join zt_user as c on a.closedBy=c.account""")
        title = '测试人员每天关闭的bug数'
        # 绘制折线图
        self.write_line(project_name, title, close_result)
        image_9 = f'{file_path}/report_res/{project_name}/{title}.png'
        # 将生成的折线图图片插入到文档对应位置
        self.save_img_to_doc(report_path, image_9, '<<test_close_days>>')
        # 计算di值
        # 取出bug等级，排除bug解决方案为 不是问题、重复bug、外部原因、无法重现 的bug
        di_result = db.select_data_list(f"""
                     select (case when r.severity=1 then 'A'
                     when r.severity=2 then 'B'
                     when r.severity=3 then 'C'
                     when r.severity=4 then 'D'
                     when r.severity=5 then 'E'
                     end)level,r.count,concat(ROUND(r.count/b.total*100, 2), '%') as res
                     from (select severity,count(*) as count from zt_bug where project={pid} and deleted='0'
                     and assignedTo='closed' and resolution not in ('external','duplicate','externalReasons','notrepro') 
                     group by severity)r, (select count(*) as total from zt_bug
                     where project={pid} and deleted='0' and assignedTo='closed' 
                     and resolution not in ('external','duplicate','externalReasons','notrepro'))b;""")
        actual = self.calcuate_di(di_result)
        self.write_string(report_path, '{{expect_DI}}', expect_di)
        self.write_string(report_path, '{{actual_DI}}', str(actual))
        diff = self.di_diff(float(expect_di), actual)
        self.write_string(report_path, '{{diff}}', str(diff))

    def write_table(self, res_doc, data, name, item=None):
        """绘制表格"""
        ori_doc = f'{file_path}/template.docx'
        # 查看文件大小，若为0，则文件无内容，取data文件夹中的测试报告模板
        if os.path.getsize(res_doc) == 0:
            tpl_doc = ori_doc
        else:
            tpl_doc = res_doc
        # 列名列表
        if item is not None:
            new_data = [['ID', '简述', '说明', '处理建议']]
            t_col = 4
        else:
            new_data = [['条目', '数量', '百分比']]
            t_col = 3
        # 将结果追加到新列表中
        for i in data:
            new_data.append(i)
        if len(new_data) == 0 or len(new_data) == 1:
            t_row = 3
        else:
            t_row = len(new_data)
        # 打开模板文件
        doc = Document(tpl_doc)
        # 根据文档中的占位符定位图片插入的位置
        target = None
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.endswith(name):
                # 把占位符去掉
                paragraph.text = paragraph.text.replace(f'{name}', '')
                run = paragraph.add_run('')
                run.add_break()
                target = paragraph
                break
        if target is not None:
            # 在文档中增加表格，并添加文字
            table = doc.add_table(rows=t_row, cols=t_col, style='Table Grid')
            for row in range(len(new_data)):
                cells = table.rows[row].cells
                for col in range(t_col):
                    cells[col].text = str(new_data[row][col])
            # 把添加的table移动到指定的位置
            self.move_table_after(table, target)
        # 保存结果文件
        doc.save(res_doc)

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def write_pie(self, project, title, data):
        """
        绘制饼图
        :param project: 项目名称
        :param title: 饼图的标题
        :param data: 用于统计的数据
        :return:
        """
        """"""
        li_1 = []
        li_2 = []
        for i in data:
            li_1.append(i[0])
            li_2.append(i[1])
        c = (
            Pie()
            .add("", [list(z) for z in zip(li_1, li_2)])
            # 自定义饼图的颜色
            # .set_colors(["blue", "green", "yellow", "red", "pink", "orange"])
            .set_global_opts(title_opts=opts.TitleOpts(title=f"{title}", pos_left="0%"))
            .set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"), position="bottom")
        )
        # 将生成的饼图转换为图片
        make_snapshot(snapshot, c.render(f"{file_path}/render.html"), f"{file_path}/report_res/{project}/{title}.png")
        # return c

    def write_line(self, project, title, data):
        """
        绘制折线图
        :param project: 项目名称
        :param title: 折线图的标题
        :param data: 用于统计的数据
        :return:
        """
        num = []
        x = []
        y1 = []
        y2 = []
        for i in data:
            if i[0] not in num:
                num.append(i[0])
            if i[1] not in x:
                x.append(i[1])
        for k in data:
            if k[0] == num[0]:
                y1.append(k[-1])
            if k[0] == num[1]:
                y2.append(k[-1])
        line = (
            Line()
            .add_xaxis(xaxis_data=x)
            .add_yaxis(series_name=num[0], y_axis=y1, is_smooth=True,
                       areastyle_opts=opts.AreaStyleOpts(opacity=0.5))
            .add_yaxis(series_name=num[-1], y_axis=y2, is_smooth=True,
                       areastyle_opts=opts.AreaStyleOpts(opacity=0.5))
            .set_global_opts(title_opts=opts.TitleOpts(title=f"{title}"))
        )
        # 将生成的折线图转换为图片
        make_snapshot(snapshot, line.render(f"{file_path}/render.html"), f"{file_path}/report_res/{project}/{title}.png")
        # return line

    def get_pareto(self, project, title, res):
        """
        绘制帕累托图
        :param project: 项目名称
        :param title: 图片的标题
        :param res: 用于统计的数据
        :return:
        """
        dict = {}    # 将结果转换为字典
        for i in res:
            dict[i[0]] = i[1]
        # 根据字典value的大小，对素进行降序排序，然后结果赋值给result
        result = sorted(dict.items(), key=lambda item: item[1], reverse=True)
        # 定义两个空列表
        typeList = []
        countList = []
        # for循环遍历列表result
        for item in result:
            # 其中元组中第一个元素就是x轴的数据，添加到列表typeList中
            typeList.append(item[0])
            # 其中元组的第二个元素就是y轴的数据，添加到列表countList中
            countList.append(item[1])
        # 定义一个累计频率的空列表percentlist
        percentlist = []
        # 定义一个计算累计频率的变量accu_percent，赋值为0
        accu_percent = 0
        # 遍历列表countList中的每个元素
        for score in countList:
            # 计算单个扣分分数，占所有扣分分数的比例，乘100，赋值给percent
            percent = score / sum(countList) * 100
            # 累计频率accu_percent，加上当前扣分分数的频率，得到当前的累计频率accu_percent
            accu_percent = accu_percent + percent
            # 使用round()函数保留两位小数，append到列表percentlist中
            percentlist.append(round(accu_percent, 2))
        # 创建一个柱状图Bar对象并赋值给变量bar
        bar = Bar()
        # 使用add_xaxis函数，传入typeList作为x轴数据
        bar.add_xaxis(xaxis_data=typeList)
        # 使用add_yaxis函数，传入countList作为y轴数据
        # 柱状图的宽度为70%，可以进行调整
        bar.add_yaxis(
            series_name="bug类型缺陷数",
            y_axis=countList,
            category_gap="70%"
        )
        # 使用全局配置项，设置x轴旋转45度，设置标题为"帕累托图分析"
        bar.set_global_opts(
            xaxis_opts=opts.AxisOpts(axislabel_opts={"rotate": 45}),
            title_opts=opts.TitleOpts(title='帕累托图分析')
        )
        # 使用extend_axis函数，参数yaxis传入坐标轴配置项opts.AxisOpts()
        bar.extend_axis(yaxis=opts.AxisOpts())
        # 创建一个折线图Line对象并赋值给变量line
        line = Line()
        # 使用add_xaxis函数，传入pointList作为x轴数据
        line.add_xaxis(xaxis_data=typeList)
        # 使用add_yaxis函数，设置yaxis_index为1，设置z_level为1
        line.add_yaxis(
            series_name="缺陷数百分比累加",
            y_axis=percentlist,
            yaxis_index=1,
            z_level=1
        )
        # 对bar使用overlap()函数，传入line，就是在柱状图的基础上叠加折线图
        bar.overlap(line)
        # 将生成的帕累托图转换为图片
        make_snapshot(snapshot, bar.render(f"{file_path}/pareto_overlap.html"), f"{file_path}/report_res/{project}/{title}.png")

    def save_img_to_doc(self, res_doc, img, name):
        """把图片保存到doc文件中的指定位置"""
        ori_doc = f'{file_path}/template.docx'
        # 查看文件大小，若无0，则文件无内容，则取data中的模板
        if os.path.getsize(res_doc) == 0:
            tpl_doc = ori_doc
        else:
            tpl_doc = res_doc
        # 打开模板文件
        doc = Document(tpl_doc)
        # 插入图片并居中
        for paragraph in doc.paragraphs:
            # 根据文档中的占位符定位图片插入的位置
            if f'{name}' in paragraph.text:
                # 把占位符去掉
                paragraph.text = paragraph.text.replace(f'{name}', '')
                run = paragraph.add_run('')
                run.add_break()
                # 添加图片并指定大小
                run.add_picture(img, width=Inches(6.2))
        # 保存结果文件
        doc.save(res_doc)

    def write_string(self, res_doc, find, rep):
        ori_doc = f'{file_path}/template.docx'
        # 查看文件大小，若无0，则文件无内容，则取data中的模板
        if os.path.getsize(res_doc) == 0:
            tpl_doc = ori_doc
        else:
            tpl_doc = res_doc
        # 打开模板文件
        doc = Document(tpl_doc)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if find in run.text:
                    # find为word中的字，rep为要替换的字
                    run.text = run.text.replace(find, rep)
        doc.save(res_doc)

    def find_max_bug(self, result):
        li = []
        for i in result:
            li.append(i[1])
        m = li[0]
        for j in li:
            if j >= m:
                m = j
                k = li.index(m)
        return result[k]

    def calcuate_di(self, result):
        """计算di值"""
        res_li = []
        for k in result:
            if k[0] == 'A':
                res_li.append(k[1]*10)
            elif k[0] == 'B':
                res_li.append(k[1]*5)
            elif k[0] == 'C':
                res_li.append(k[1]*1)
            elif k[0] == 'D':
                res_li.append(k[1] * 0.5)
            elif k[0] == 'E':
                res_li.append(k[1]*0.1)
        return sum(res_li)

    def di_diff(self, expect, actual):
        """计算项目实际的di值与预估的di值之间的差别"""
        if actual > expect:
            diff = '比预期数量多。'
        elif actual < expect:
            diff = '比预期数量少。'
        else:
            diff = '与预期数量一致。'
        return diff
# #

# if __name__ == '__main__':
#     pytest.main([__file__])

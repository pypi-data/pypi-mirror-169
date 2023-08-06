import os
import numpy as np
import pandas as pd
import logging
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx import Document
from docx.shared import Inches
from datetime import datetime
from ngdataqualitycheck.helper import *
from ngdataqualitycheck.plot_builder import *
from ngdataqualitycheck.data_stats import *
from ngdataqualitycheck.global_params import *
from ngdataqualitycheck.make_directory import *
import warnings
warnings.filterwarnings('ignore')


def generate_report(filename, enable_logging=False):
    """
    For generating the data quality report.

    :param filename: Name of the file for which data quality report has to be generated
    :type filename: str
    :param enable_logging: To generate logfiles, defaults to False
    "type enable_logging: bool, optional
    :return: Data quality report.
    :rtype: docx, csv
    """
    name_of_file = filename.split('.')[0]
    plots_folder = plot_folder

    if enable_logging:
        # Instantiating MakeDirectory class
        md = MakeDirectory(
            dir_structure=generate_directory['dir_structure'],
            basefolder=generate_directory['basefolder'])

        # Creating logging directory
        logfile_name = md.create_logging_file(foldername='logging')
        # Creating directory for saving docs
        docs_folder = md.create_directory(foldername='output\\docs')
        # Creating directory for saving csv
        csv_folder = md.create_directory(foldername='output\\csv')

        logging.basicConfig(format='%(asctime)s : %(message)s',
                            filename=logfile_name,
                            level=logging.INFO)

        app_logger = logging.getLogger()

        app_logger.info("<< Initializing Data Qualtiy Report generation >>")

        # GENERATING VALUES
        # Fetching date and time
        datenow, timenow = get_datetime()

        try:
            app_logger.info("Loading data")
            # Loading data
            df = load_data(filename=filename)
        except Exception:
            logging.exception("Failed at loading data")
            raise Exception()

        try:
            app_logger.info("Obtaining data shape and dtypes")
            # Getting shape and dtypes
            df_shape, typeslist = get_data_info(df=df)
        except Exception:
            logging.exception("Failed at obtaining data shape and dtypes")
            raise Exception()

        try:
            app_logger.info("Obtaining data for empty features")
            # Getting Empty's data
            emptys_count, emptys_list = get_emptys(df=df)
        except Exception:
            logging.exception("Failed at obtaining data for empty features")
            raise Exception()

        try:
            app_logger.info("Obtaining data for features with 0's")
            # Getting Zero's data
            _, zeros_list = get_zeros(df=df)
        except Exception:
            logging.exception("Failed at obtaining data for features with 0's")
            raise Exception()

        try:
            app_logger.info("Obtaining data for features with negative values")
            # Getting Negative's data
            _, negatives_list = get_negatives(df=df)
        except Exception:
            logging.exception(
                "Failed at obtaining data for features with negative values")
            raise Exception()

        try:
            app_logger.info(
                "Obtaining data for features with duplicate values")
            # Getting Duplicate's data
            num_dups, _, dup_cols = check_duplicates(df=df)
        except Exception:
            logging.exception(
                "Failed at obtaining data for features with duplicate values")
            raise Exception()

        try:
            app_logger.info("Obtaining data for features with missing values")
            # Getting Missing's data
            num_miss, miss_cols = check_missing(df=df)
        except Exception:
            logging.exception(
                "Failed at obtaining data for features with missing values")
            raise Exception()

        try:
            app_logger.info("Obtaining data for features with constant values")
            # Getting Constants's data
            num_const, const_cols = check_constants(df=df)
        except Exception:
            logging.exception(
                "Failed at obtaining data for features with constant values")
            raise Exception()

        try:
            app_logger.info("Obtaining stats for text features")
            # Getting Text data
            df, obj_cols, txt_coldict = get_text_stats(df=df)
        except Exception:
            logging.exception("Failed at obtaining stats for text features")
            raise Exception()

        try:
            app_logger.info("Obtaining stats for numerical features")
            # Getting Numerical data
            df, num_cols, discrete_cols, continuous_cols, num_coldict = get_num_stats(
                df=df)
        except Exception:
            logging.exception(
                "Failed at obtaining stats for numerical features")
            raise Exception()

        try:
            app_logger.info("Obtaining stats for categorical features")
            # Getting Categorical data
            df, cat_cols, cat_coldict = get_cat_stats(df=df)
        except Exception:
            logging.exception("Failed at obtaining stats categorical features")
            raise Exception()

        try:
            app_logger.info("Obtaining data for correlated features")
            # Getting Correlation data
            corlist = get_correlation(df=df, thresh=CORR_THRESHOLD)
        except Exception:
            logging.exception(
                "Failed at obtaining data for correlated features")
            raise Exception()

        try:
            app_logger.info("Obtaining stats for datetime features")
            # Getting Datetime data
            df, time_cols, time_coldict = get_time_stats(df=df)
        except Exception:
            logging.exception(
                "Failed at obtaining stats for datetime features")
            raise Exception()

        try:
            app_logger.info("Generating warnings")
            # Generating Warnings
            warnings_dict = generate_warnings(
                df=df,
                txt_coldict=txt_coldict,
                cat_coldict=cat_coldict,
                num_coldict=num_coldict,
                time_coldict=time_coldict)
        except Exception:
            logging.exception("Failed at generating warnings")
            raise Exception()

        try:
            app_logger.info("Plotting numerical data")
            # Plotting Numerical data
            for col in continuous_cols:
                plot_numerical(
                    df=df,
                    name_of_file=name_of_file,
                    col=col,
                    timenow=timenow)
        except Exception:
            logging.exception("Failed at plotting numerical data")
            raise Exception()

        try:
            app_logger.info("Plotting discrete data")
            # Plotting Discrete data
            for col in discrete_cols:
                plot_discrete(
                    df=df,
                    name_of_file=name_of_file,
                    col=col,
                    timenow=timenow)
        except Exception:
            logging.exception("Failed at plotting discrete data")
            raise Exception()

        try:
            app_logger.info("Plotting categorical data")
            # Plotting Categorical data
            for col in cat_cols:
                plot_discrete(
                    df=df,
                    name_of_file=name_of_file,
                    col=col,
                    timenow=timenow)
        except Exception:
            logging.exception("Failed at plotting categorical data")
            raise Exception()

        try:
            app_logger.info("Plotting text data")
            # Plotting Text data
            plot_text(
                df=df,
                name_of_file=name_of_file,
                text_cols_list=obj_cols,
                timenow=timenow)
        except Exception:
            logging.exception("Failed at plotting text data")
            raise Exception()

        try:
            app_logger.info("Plotting missing values")
            # Generating missing values plots
            miss_bar_name, miss_hmap_name = plot_missing(
                df=df,
                name_of_file=name_of_file,
                timenow=timenow)
        except Exception:
            logging.exception("Failed at plotting missing values")
            raise Exception()

        # <<<   CREATING DOCUMENT   >>>

        # Functions for adding page numbers

        def create_element(name):
            """
            For creating a document element

            :param name: Field character
            :type name: str
            :return: element
            :rtype: str
            """
            return OxmlElement(name)

        def create_attribute(element, name, value):
            """
            For creating a page attribute

            :param element: Element
            :type element: str
            :param name: Field character
            :type name: str
            :param value: Value
            :type value: str
            """
            element.set(ns.qn(name), value)

        def add_page_number(run):
            """
            For creating page number

            :param run: Run
            :type run: str
            """
            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        try:
            app_logger.info("Generating Data Quality Report document")
            # Initializing document object
            document = Document()

            # Adding Title
            for i in range(5):
                document.add_paragraph()

            title = document.add_heading(text='NgDataQualityReport', level=0)
            title_format = title.paragraph_format
            title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adding Sub-Title
            for i in range(3):
                document.add_paragraph()

            sub_title = document.add_heading(
                text=f'Ngenux Solutions Pvt. Ltd.', level=1)
            sub_title_format = sub_title.paragraph_format
            sub_title_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Adding Sub-Title date
            date_text = datetime.now().strftime("%B %d, %Y")
            sub_title = document.add_heading(text=date_text, level=3)
            sub_title_format = sub_title.paragraph_format
            sub_title_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Copyright footer
            footer1 = document.sections[0].footer
            footer1_para = footer1.paragraphs[0]
            footer1_para.alignment = 2
            footer1_run = footer1_para.add_run()
            year_now = datetime.now().year
            footer1_run.text = f"© Copyright {year_now}, Ngenux Solutions Pvt. Ltd.\t\t"

            # Creating page number
            page_num_footer = document.sections[0].footer
            page_num_footer_para = page_num_footer.paragraphs[0]
            page_num_footer_para.alignment = 2
            page_num_footer_run = page_num_footer_para.add_run()
            add_page_number(page_num_footer_run)

            # Adding a pagebreak
            document.add_page_break()

            # <<<      GLOBAL STATS     >>>

            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Global Statistics for the data', level=0)

            # generating variables
            file_type = filename.split('.')[-1]
            file_encoding = get_encoding(filename=filename)
            total_samples = df_shape[0]
            total_features = df_shape[1]
            total_empty_columns = emptys_count
            per_samples_with_nan = str(
                np.round((num_miss/total_samples)*100, 2))+"%"
            total_missing_cols = len(miss_cols)
            total_duplicate_rows = num_dups
            total_duplicate_cols = len(dup_cols)
            total_const = num_const

            # Creating records for table
            global_records = (
                ('File Name', filename),
                ('File type', file_type),
                ('File encoding', file_encoding),
                ('Number of Rows', total_samples),
                ('Number of Columns', total_features),
                ('Number of Columns without any values', total_empty_columns),
                ('Percentage of Rows having Missing values', per_samples_with_nan),
                ('Number of Columns having Missing values', total_missing_cols),
                ('Number of Duplicate Rows', total_duplicate_rows),
                ('Number of Duplicate Columns', total_duplicate_cols),
                ('Number of Columns with no Variance', total_const),
            )

            # Creating table
            table = document.add_table(
                rows=1, cols=2, style='LightShading-Accent1')
            # Creating table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Inference'
            hdr_cells[1].text = 'Values'
            # Adding data into table rows
            for i, j in global_records:
                row_cells = table.add_row().cells
                row_cells[0].text = i
                row_cells[1].text = str(j)

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding a pagebreak
            document.add_page_break()

            # <<<      PROFILE SCHEMA     >>>

            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Profile Schema', level=0)

            # Creating records
            profile_schema_records = typeslist

            # Creating table
            table = document.add_table(
                rows=1, cols=3, style='LightShading-Accent1')

            # Creating table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No'
            hdr_cells[1].text = 'Column'
            hdr_cells[2].text = 'Datatype as per source'
            # Adding data into table rows
            for i, j, k in profile_schema_records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = j
                row_cells[2].text = k

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding a pagebreak
            document.add_page_break()

            # <<<      MISSING VALUE PLOTS     >>>

            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Missing Value Plots', level=0)

            # Bar Chart
            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Bar Chart', level=1)

            # Generating dictionary of missing values
            missing_dict = {}
            for col in df:
                missing_dict[col] = np.round(100*(df[col].isnull().mean()), 2)
            missing_dict = {
                k: v for k, v in sorted(missing_dict.items(),
                                        key=lambda item: item[1],
                                        reverse=True)}
            missing_dict = {k: v for k, v in missing_dict.items() if v != 0}

            # Adding a paragraph
            document.add_paragraph

            # Adding Missing bar plot
            document.add_picture(miss_bar_name, width=Inches(6))

            # Adding a pagebreak
            document.add_page_break()

            # Heatmap
            # Adding a paragraph
            p = document.add_paragraph()

            # Adding heading
            document.add_heading('Heatmap', level=1)

            # Adding a picture
            document.add_picture(miss_hmap_name, width=Inches(6))

            # Adding a pagebreak
            document.add_page_break()

            # <<<      CORRELATION PLOT     >>>

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding heading
            document.add_heading('Correlation Plot', level=0)

            # Generating Correlation Plots
            if len(corlist) > 0:
                # Generating Correlation plots
                corr_plot_name = plot_correlation(
                    df=df,
                    name_of_file=name_of_file,
                    corlist=corlist,
                    timenow=timenow)

                # Adding a paragraph
                p = document.add_paragraph()

                # Adding Correlation plots
                document.add_picture(corr_plot_name, width=Inches(6))
            else:
                document.add_paragraph()
                document.add_paragraph(
                    'Unable to generate plot', style='Intense Quote')
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(
                    f'Correlated Variables above "{CORR_THRESHOLD}" threshold not found.')
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adding a pagebreak
            document.add_page_break()

            # <<<      SUMMARY     >>>

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding heading
            document.add_heading('Profile Summary', level=0)

            # Fetching the list of plots
            plots_list = []

            for item in os.listdir(plots_folder):
                if item.split('.')[-1] == PLOT_FORMAT[0]:
                    plots_list.append(item)

            for i in range(total_features):
                # Obtaining source dictionary
                if df.columns[i] in time_cols:
                    src_dict = time_coldict[df.columns[i]]
                elif df.columns[i] in obj_cols:
                    src_dict = txt_coldict[df.columns[i]]
                elif df.columns[i] in cat_cols:
                    src_dict = cat_coldict[df.columns[i]]
                elif df.columns[i] in num_cols:
                    src_dict = num_coldict[df.columns[i]]

                # Generating records
                records = []
                for k, v in src_dict.items():
                    records.append((k, v))
                records = tuple(records)

                # Adding heading
                document.add_heading(text=f'Column {i+1}', level=2)

                # Creating table
                table = document.add_table(
                    rows=1, cols=2, style='LightShading-Accent1')

                # Creating table headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Inference'
                hdr_cells[1].text = 'Values'

                # Adding data into table rows
                for rec1, rec2 in records:
                    row_cells = table.add_row().cells
                    row_cells[0].text = rec1
                    row_cells[1].text = str(rec2)

                # fetching plots
                col_plot = []
                for plot in plots_list:
                    if plot == f'{name_of_file}_{df.columns[i]}_{timenow}.png':
                        col_plot.append(plot)

                # Adding a paragraph
                document.add_paragraph()

                # Adding a picture
                if len(col_plot) > 0:
                    pic_name = os.path.join(plots_folder, col_plot[0])
                    document.add_picture(pic_name, width=Inches(4))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    document.add_paragraph()
                    document.add_paragraph(
                        'Unable to generate plot', style='Intense Quote')
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adding a pagebreak
                document.add_page_break()

            # creating filename
            op_file = f'{docs_folder}\\{name_of_file}_data_quality_report_{datenow}_{timenow}.docx'

            # Saving the document
            document.save(op_file)
        except Exception:
            logging.exception(
                "Failed at generating Data Quality Report document")
            raise Exception()

        # <<<      GENERATING CSV     >>>

        # <<<      GLOBAL STATISTICS     >>>

        try:
            app_logger.info("Generating Global Statistics csv")
            global_records_dict = dict(global_records)

            global_statistics_dict = {
                'Attributes': list(global_records_dict.keys()),
                'Values': list(global_records_dict.values())
            }

            global_statistics = pd.DataFrame.from_dict(global_statistics_dict)

            global_statistics_filename = f'{csv_folder}\\{name_of_file}_global_statistics_{datenow}_{timenow}.csv'

            # Saving csv
            global_statistics.to_csv(global_statistics_filename, index=False)
        except Exception:
            logging.exception("Failed at generating Global Statistics csv")
            raise Exception()

        # <<<      PROFILE SCHEMA     >>>

        try:
            app_logger.info("Generating Profile Schema csv")
            variable = [i[1] for i in typeslist]
            variable_type = [i[2] for i in typeslist]

            profile_schema_dict = {
                'Variable_Name': variable,
                'Data_Type': variable_type
            }

            profile_schema = pd.DataFrame.from_dict(profile_schema_dict)

            profile_schema_filename = f'{csv_folder}\\{name_of_file}_profile_schema_{datenow}_{timenow}.csv'

            # Saving csv
            profile_schema.to_csv(profile_schema_filename, index=False)
        except Exception:
            logging.exception("Failed at generating Profile Schema csv")
            raise Exception()

        # <<<      WARNINGS     >>>

        try:
            app_logger.info("Generating Warnings csv")
            
            warnings_df = pd.DataFrame()

            warnings_df['Variable_Name'] = variable
            warnings_df['Has_Missing_Values'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['missing_values']),
                1, 0)
            warnings_df['Is_Imbalanced'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['imbalance']),
                1, 0)
            warnings_df['Has_Negative_Mean'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['neg_mean']),
                1, 0)
            warnings_df['Is_Left_Skewed'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['left_skew']),
                1, 0)
            warnings_df['Is_Right_Skewed'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['right_skew']),
                1, 0)
            warnings_df['Has_Outliers'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['outlier']),
                1, 0)
            warnings_df['Is_Empty'] = np.where(
                warnings_df['Variable_Name'].isin(emptys_list),
                1, 0)
            warnings_df['Has_Zero'] = np.where(
                warnings_df['Variable_Name'].isin(zeros_list),
                1, 0)
            warnings_df['Has_Negatives'] = np.where(
                warnings_df['Variable_Name'].isin(negatives_list),
                1, 0)
            warnings_df['Is_Constant'] = np.where(
                warnings_df['Variable_Name'].isin(const_cols),
                1, 0)
            warnings_df['Is_Discrete'] = np.where(
                warnings_df['Variable_Name'].isin(discrete_cols),
                1, 0)
            warnings_df['Is_Continuous'] = np.where(
                warnings_df['Variable_Name'].isin(continuous_cols),
                1, 0)
            warnings_df['Is_Categorical'] = np.where(
                warnings_df['Variable_Name'].isin(cat_cols),
                1, 0)
            warnings_df['Is_Datetime'] = np.where(
                warnings_df['Variable_Name'].isin(time_cols),
                1, 0)

            warnings_filename = f'{csv_folder}\\{name_of_file}_warnings_{datenow}_{timenow}.csv'

            # Saving csv
            warnings_df.to_csv(warnings_filename, index=False)
        except Exception:
            logging.exception("Failed at generating Warnings csv")
            raise Exception()


    else:
        # Instantiating MakeDirectory class
        md = MakeDirectory(
            dir_structure=generate_directory['dir_structure'],
            basefolder=generate_directory['basefolder'])

        # Creating directory for saving docs
        docs_folder = md.create_directory(foldername='output\\docs')
        # Creating directory for saving csv
        csv_folder = md.create_directory(foldername='output\\csv')

        # GENERATING VALUES
        # Fetching date and time
        datenow, timenow = get_datetime()

        try:
            # Loading data
            df = load_data(filename=filename)
        except Exception:
            raise Exception()

        try:
            # Getting shape and dtypes
            df_shape, typeslist = get_data_info(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Empty's data
            emptys_count, emptys_list = get_emptys(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Zero's data
            _, zeros_list = get_zeros(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Negative's data
            _, negatives_list = get_negatives(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Duplicate's data
            num_dups, _, dup_cols = check_duplicates(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Missing's data
            num_miss, miss_cols = check_missing(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Constants's data
            num_const, const_cols = check_constants(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Text data
            df, obj_cols, txt_coldict = get_text_stats(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Numerical data
            df, num_cols, discrete_cols, continuous_cols, num_coldict = get_num_stats(
                df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Categorical data
            df, cat_cols, cat_coldict = get_cat_stats(df=df)
        except Exception:
            raise Exception()

        try:
            # Getting Correlation data
            corlist = get_correlation(df=df, thresh=CORR_THRESHOLD)
        except Exception:
            raise Exception()

        try:
            # Getting Datetime data
            df, time_cols, time_coldict = get_time_stats(df=df)
        except Exception:
            raise Exception()

        try:
            # Generating Warnings
            warnings_dict = generate_warnings(
                df=df,
                txt_coldict=txt_coldict,
                cat_coldict=cat_coldict,
                num_coldict=num_coldict,
                time_coldict=time_coldict)
        except Exception:
            raise Exception()

        try:
            # Plotting Numerical data
            for col in continuous_cols:
                plot_numerical(
                    df=df,
                    name_of_file=name_of_file,
                    col=col,
                    timenow=timenow)
        except Exception:
            raise Exception()

        try:
            # Plotting Discrete data
            for col in discrete_cols:
                plot_discrete(
                    df=df,
                    name_of_file=name_of_file,
                    col=col,
                    timenow=timenow)
        except Exception:
            raise Exception()

        try:
            # Plotting Categorical data
            for col in cat_cols:
                plot_discrete(
                    df=df,
                    name_of_file=name_of_file,
                    col=col,
                    timenow=timenow)
        except Exception:
            raise Exception()

        try:
            # Plotting Text data
            plot_text(
                df=df,
                name_of_file=name_of_file,
                text_cols_list=obj_cols,
                timenow=timenow)
        except Exception:
            raise Exception()

        try:
            # Generating missing values plots
            miss_bar_name, miss_hmap_name = plot_missing(
                df=df, name_of_file=name_of_file, timenow=timenow)
        except Exception:
            raise Exception()

        # <<<   CREATING DOCUMENT   >>>

        # Functions for adding page numbers

        def create_element(name):
            """
            For creating a document element

            :param name: Field character
            :type name: str
            :return: element
            :rtype: str
            """
            return OxmlElement(name)

        def create_attribute(element, name, value):
            """
            For creating a page attribute

            :param element: Element
            :type element: str
            :param name: Field character
            :type name: str
            :param value: Value
            :type value: str
            """
            element.set(ns.qn(name), value)

        def add_page_number(run):
            """
            For creating page number

            :param run: Run
            :type run: str
            """
            fldChar1 = create_element('w:fldChar')
            create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = create_element('w:instrText')
            create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = create_element('w:fldChar')
            create_attribute(fldChar2, 'w:fldCharType', 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        try:
            # Initializing document object
            document = Document()

            # Adding Title
            for i in range(5):
                document.add_paragraph()

            title = document.add_heading(text='NgDataQualityReport', level=0)
            title_format = title.paragraph_format
            title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adding Sub-Title
            for i in range(3):
                document.add_paragraph()

            sub_title = document.add_heading(
                text=f'Ngenux Solutions Pvt. Ltd.', level=1)
            sub_title_format = sub_title.paragraph_format
            sub_title_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Adding Sub-Title date
            date_text = datetime.now().strftime("%B %d, %Y")
            sub_title = document.add_heading(text=date_text, level=3)
            sub_title_format = sub_title.paragraph_format
            sub_title_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Copyright footer
            footer1 = document.sections[0].footer
            footer1_para = footer1.paragraphs[0]
            footer1_para.alignment = 2
            footer1_run = footer1_para.add_run()
            year_now = datetime.now().year
            footer1_run.text = f"© Copyright {year_now}, Ngenux Solutions Pvt. Ltd.\t\t"

            # Creating page number
            page_num_footer = document.sections[0].footer
            page_num_footer_para = page_num_footer.paragraphs[0]
            page_num_footer_para.alignment = 2
            page_num_footer_run = page_num_footer_para.add_run()
            add_page_number(page_num_footer_run)

            # Adding a pagebreak
            document.add_page_break()

            # <<<      GLOBAL STATS     >>>

            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Global Statistics for the data', level=0)

            # generating variables
            file_type = filename.split('.')[-1]
            file_encoding = get_encoding(filename=filename)
            total_samples = df_shape[0]
            total_features = df_shape[1]
            total_empty_columns = emptys_count
            per_samples_with_nan = str(
                np.round((num_miss/total_samples)*100, 2))+"%"
            total_missing_cols = len(miss_cols)
            total_duplicate_rows = num_dups
            total_duplicate_cols = len(dup_cols)
            total_const = num_const

            # Creating records for table
            global_records = (
                ('File Name', filename),
                ('File type', file_type),
                ('File encoding', file_encoding),
                ('Number of Rows', total_samples),
                ('Number of Columns', total_features),
                ('Number of Columns without any values', total_empty_columns),
                ('Percentage of Rows having Missing values', per_samples_with_nan),
                ('Number of Columns having Missing values', total_missing_cols),
                ('Number of Duplicate Rows', total_duplicate_rows),
                ('Number of Duplicate Columns', total_duplicate_cols),
                ('Number of Columns with no Variance', total_const),
            )

            # Creating table
            table = document.add_table(
                rows=1, cols=2, style='LightShading-Accent1')
            # Creating table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Inference'
            hdr_cells[1].text = 'Values'
            # Adding data into table rows
            for i, j in global_records:
                row_cells = table.add_row().cells
                row_cells[0].text = i
                row_cells[1].text = str(j)

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding a pagebreak
            document.add_page_break()

            # <<<      PROFILE SCHEMA     >>>

            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Profile Schema', level=0)

            # Creating records
            profile_schema_records = typeslist

            # Creating table
            table = document.add_table(
                rows=1, cols=3, style='LightShading-Accent1')

            # Creating table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No'
            hdr_cells[1].text = 'Column'
            hdr_cells[2].text = 'Datatype as per source'
            # Adding data into table rows
            for i, j, k in profile_schema_records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = j
                row_cells[2].text = k

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding a pagebreak
            document.add_page_break()

            # <<<      MISSING VALUE PLOTS     >>>

            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Missing Value Plots', level=0)

            # Bar Chart
            # Adding a paragraph
            p = document.add_paragraph()
            # Adding heading
            document.add_heading('Bar Chart', level=1)

            # Generating dictionary of missing values
            missing_dict = {}
            for col in df:
                missing_dict[col] = np.round(100*(df[col].isnull().mean()), 2)
            missing_dict = {
                k: v for k, v in sorted(missing_dict.items(),
                                        key=lambda item: item[1],
                                        reverse=True)}
            missing_dict = {k: v for k, v in missing_dict.items() if v != 0}

            # Adding a paragraph
            document.add_paragraph

            # Adding Missing bar plot
            document.add_picture(miss_bar_name, width=Inches(6))

            # Adding a pagebreak
            document.add_page_break()

            # Heatmap
            # Adding a paragraph
            p = document.add_paragraph()

            # Adding heading
            document.add_heading('Heatmap', level=1)

            # Adding a picture
            document.add_picture(miss_hmap_name, width=Inches(6))

            # Adding a pagebreak
            document.add_page_break()

            # <<<      CORRELATION PLOT     >>>

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding heading
            document.add_heading('Correlation Plot', level=0)

            # Generating Correlation Plots
            if len(corlist) > 0:
                # Generating Correlation plots
                corr_plot_name = plot_correlation(
                    df=df, name_of_file=name_of_file, corlist=corlist, timenow=timenow)

                # Adding a paragraph
                p = document.add_paragraph()

                # Adding Correlation plots
                document.add_picture(corr_plot_name, width=Inches(6))
            else:
                document.add_paragraph()
                document.add_paragraph(
                    'Unable to generate plot', style='Intense Quote')
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(
                    f'Correlated Variables above "{CORR_THRESHOLD}" threshold not found.')
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adding a pagebreak
            document.add_page_break()

            # <<<      SUMMARY     >>>

            # Adding a paragraph
            p = document.add_paragraph()

            # Adding heading
            document.add_heading('Profile Summary', level=0)

            # Fetching the list of plots
            plots_list = []

            for item in os.listdir(plots_folder):
                if item.split('.')[-1] == PLOT_FORMAT[0]:
                    plots_list.append(item)

            for i in range(total_features):
                # Obtaining source dictionary
                if df.columns[i] in time_cols:
                    src_dict = time_coldict[df.columns[i]]
                elif df.columns[i] in obj_cols:
                    src_dict = txt_coldict[df.columns[i]]
                elif df.columns[i] in cat_cols:
                    src_dict = cat_coldict[df.columns[i]]
                elif df.columns[i] in num_cols:
                    src_dict = num_coldict[df.columns[i]]

                # Generating records
                records = []
                for k, v in src_dict.items():
                    records.append((k, v))
                records = tuple(records)

                # Adding heading
                document.add_heading(text=f'Column {i+1}', level=2)

                # Creating table
                table = document.add_table(
                    rows=1, cols=2, style='LightShading-Accent1')

                # Creating table headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Inference'
                hdr_cells[1].text = 'Values'

                # Adding data into table rows
                for rec1, rec2 in records:
                    row_cells = table.add_row().cells
                    row_cells[0].text = rec1
                    row_cells[1].text = str(rec2)

                # fetching plots
                col_plot = []
                for plot in plots_list:
                    if plot == f'{name_of_file}_{df.columns[i]}_{timenow}.png':
                        col_plot.append(plot)

                # Adding a paragraph
                document.add_paragraph()

                # Adding a picture
                if len(col_plot) > 0:
                    pic_name = os.path.join(plots_folder, col_plot[0])
                    document.add_picture(pic_name, width=Inches(4))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    document.add_paragraph()
                    document.add_paragraph(
                        'Unable to generate plot', style='Intense Quote')
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adding a pagebreak
                document.add_page_break()

            for i in range(10):
                document.add_paragraph()

            # creating filename
            op_file = f'{docs_folder}\\{name_of_file}_data_quality_report_{datenow}_{timenow}.docx'

            # Saving the document
            document.save(op_file)
        except Exception:
            raise Exception()

        # <<<      GENERATING CSV     >>>

        # <<<      GLOBAL STATISTICS     >>>

        try:
            global_records_dict = dict(global_records)

            global_statistics_dict = {
                'Attributes': list(global_records_dict.keys()),
                'Values': list(global_records_dict.values())
            }

            global_statistics = pd.DataFrame.from_dict(global_statistics_dict)

            global_statistics_filename = f'{csv_folder}\\{name_of_file}_global_statistics_{datenow}_{timenow}.csv'

            # Saving csv
            global_statistics.to_csv(global_statistics_filename, index=False)
        except Exception:
            raise Exception()

        # <<<      PROFILE SCHEMA     >>>

        try:
            variable = [i[1] for i in typeslist]
            variable_type = [i[2] for i in typeslist]

            profile_schema_dict = {
                'Variable_Name': variable,
                'Data_Type': variable_type
            }

            profile_schema = pd.DataFrame.from_dict(profile_schema_dict)

            profile_schema_filename = f'{csv_folder}\\{name_of_file}_profile_schema_{datenow}_{timenow}.csv'

            # Saving csv
            profile_schema.to_csv(profile_schema_filename, index=False)
        except Exception:
            raise Exception()

        # <<<      WARNINGS     >>>

        try:
            warnings_df = pd.DataFrame()

            warnings_df['Variable_Name'] = variable
            warnings_df['Has_Missing_Values'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['missing_values']),
                1, 0)
            warnings_df['Is_Imbalanced'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['imbalance']),
                1, 0)
            warnings_df['Has_Negative_Mean'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['neg_mean']),
                1, 0)
            warnings_df['Is_Left_Skewed'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['left_skew']),
                1, 0)
            warnings_df['Is_Right_Skewed'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['right_skew']),
                1, 0)
            warnings_df['Has_Outliers'] = np.where(
                warnings_df['Variable_Name'].isin(warnings_dict['outlier']),
                1, 0)
            warnings_df['Is_Empty'] = np.where(
                warnings_df['Variable_Name'].isin(emptys_list),
                1, 0)
            warnings_df['Has_Zero'] = np.where(
                warnings_df['Variable_Name'].isin(zeros_list),
                1, 0)
            warnings_df['Has_Negatives'] = np.where(
                warnings_df['Variable_Name'].isin(negatives_list),
                1, 0)
            warnings_df['Is_Constant'] = np.where(
                warnings_df['Variable_Name'].isin(const_cols),
                1, 0)
            warnings_df['Is_Discrete'] = np.where(
                warnings_df['Variable_Name'].isin(discrete_cols),
                1, 0)
            warnings_df['Is_Continuous'] = np.where(
                warnings_df['Variable_Name'].isin(continuous_cols),
                1, 0)
            warnings_df['Is_Categorical'] = np.where(
                warnings_df['Variable_Name'].isin(cat_cols),
                1, 0)
            warnings_df['Is_Datetime'] = np.where(
                warnings_df['Variable_Name'].isin(time_cols),
                1, 0)

            warnings_filename = f'{csv_folder}\\{name_of_file}_warnings_{datenow}_{timenow}.csv'

            # Saving csv
            warnings_df.to_csv(warnings_filename, index=False)
        except Exception:
            raise Exception()
